import docx

# d = docx.Document("demo.docx")

# print(d.paragraphs[0].text)  # Prints the first paragraph text
# print(d.paragraphs[1].text)
# p = d.paragraphs[1]
# print(p.runs[0].text)  # Paragraphs are split in runs depending on bold and
# italics
# print(p.runs[1].text)
# print(p.runs[2].text)
# print(p.runs[3].text)
#
# print(p.runs[1].bold)  # Returns true if the run is bold

# p.runs[3].underline = True  # Changed the 4th run to underlined
# p.runs[3].text = "italic and underlined."  # Changed the text too
# d.save("demo2.docx")  # Saved as demo2.docx

# print(p.style)
# p.style = "Title"  # Changed this paragraph to be part of title
# d.save("demo3.docx")

d = docx.Document()  # Creates a new document if no string is entered
d.add_paragraph("Hello this is a paragraph.")  # Adding a new paragraph
d.add_paragraph("This is another paragraph.")
# d.save("demo4.docx")
p = d.paragraphs[0]
p.add_run("This is a new run.")  # Added new run to first paragraph
p.runs[1].bold = True  # Made this run bold
d.save("demo5.docx")
